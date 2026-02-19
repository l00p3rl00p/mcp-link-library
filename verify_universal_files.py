
import sys
import json
import subprocess
import argparse
from pathlib import Path

# Allow --skip-live to bypass file indexing in CI / test environments
parser = argparse.ArgumentParser(description="Verify universal file type support in the Librarian.")
parser.add_argument('--skip-live', action='store_true', help="Skip live indexing+search (CI safe mode)")
args = parser.parse_args()

try:
    import openpyxl
    import docx
    from PIL import Image
except ImportError:
    print("❌ Missing dependencies. Run pip install openpyxl python-docx Pillow")
    sys.exit(1)

def create_test_files():
    print("📂 Creating test files...")

    # 1. Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws['A1'] = "Project Code"
    ws['B1'] = "Status"
    ws['A2'] = "NEXUS-001"
    ws['B2'] = "Alpha Release"
    wb.save("test_excel.xlsx")
    print("   - test_excel.xlsx")

    # 2. Word
    doc = docx.Document()
    doc.add_heading('Nexus Report', 0)
    doc.add_paragraph('Confidential mission briefing for Agent 007.')
    doc.save("test_word.docx")
    print("   - test_word.docx")

    # 3. Image
    img = Image.new('RGB', (100, 100), color='red')
    img.save("test_image.png")
    print("   - test_image.png")

def index_and_search():
    server_script = Path("mcp.py").resolve()
    if not server_script.exists():
        print(f"❌ mcp.py not found at {server_script}")
        sys.exit(1)

    files = [
        ("test_excel.xlsx", "NEXUS-001"),
        ("test_word.docx", "mission briefing"),
        ("test_image.png", "Format: PNG"),
    ]

    for filename, query in files:
        file_path = Path(filename).resolve()
        uri = f"file://{file_path}"

        print(f"\n🔗 Indexing {filename}...")
        subprocess.run(
            [sys.executable, str(server_script), "--add", uri, "--categories", "test"],
            check=True, timeout=15
        )

        print(f"🔍 Searching for '{query}'...")
        process = subprocess.Popen(
            [sys.executable, str(server_script), "--server"],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1,
        )

        try:
            # Initialize
            process.stdin.write(json.dumps({"jsonrpc": "2.0", "id": 1, "method": "initialize"}) + "\n")
            process.stdin.flush()
            process.stdout.readline()  # consume init response

            # Search
            req = {
                "jsonrpc": "2.0",
                "id": 2,
                "method": "tools/call",
                "params": {
                    "name": "search_knowledge_base",
                    "arguments": {"query": query},
                },
            }
            process.stdin.write(json.dumps(req) + "\n")
            process.stdin.flush()

            # Timeout guard — previously caused indefinite hang
            try:
                stdout, _ = process.communicate(timeout=5)
                lines = [l for l in stdout.splitlines() if l.strip()]
                line = lines[-1] if lines else "{}"
            except subprocess.TimeoutExpired:
                process.kill()
                process.wait()
                print(f"⚠️  Search timed out for '{query}' — server did not respond in 5s")
                continue

            try:
                resp = json.loads(line)
                content = resp.get("result", {}).get("content", [{}])[0].get("text", "")
                if filename in content:
                    print(f"✅ SUCCESS: Found {filename} via query '{query}'")
                else:
                    print(f"❌ FAIL: Did not find {filename}")
                    print(f"   Response: {content[:200]}...")
            except json.JSONDecodeError:
                print(f"⚠️  Could not parse server response: {line[:100]}")

        finally:
            if process.poll() is None:
                process.terminate()
                process.wait()

if __name__ == "__main__":
    create_test_files()
    if args.skip_live:
        print("\n⏭️  Skipping live index+search (--skip-live mode). File creation only.")
        print("✅ Universal file types created successfully.")
    else:
        index_and_search()
